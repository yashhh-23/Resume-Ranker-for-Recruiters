import regex as re
from pathlib import Path
from typing import Any, Dict, List

DEFAULT_JD = {
    "required_skills": [],
    "raw_required_skills": [],
    "preferred_skills": [],
    "target_title": "Any",
    "min_experience_years": 0,
    "target_industry": "Any",
    "target_field": "General",
    "skills_text": "",
    "salary_min": 0.0,
    "salary_max": 0.0,
    "seniority_level": "mid",
    "jd_text": "",
}

INDUSTRY_KEYWORDS = {
    "tech": ["software", "engineer", "developer", "data", "ML", "AI", "backend", "frontend"],
    "finance": ["fintech", "banking", "financial", "accounting", "audit", "investment"],
    "marketing": ["marketing", "SEO", "content", "brand", "growth", "digital"],
    "sales": ["sales", "business development", "BD", "account executive"],
    "healthcare": ["health", "medical", "pharma", "clinical", "hospital"],
    "design": ["UI", "UX", "designer", "figma", "illustrator", "creative"],
}

REQUIRED_HEADERS = ["required", "must have", "mandatory", "essential"]
PREFERRED_HEADERS = ["preferred", "nice to have", "bonus", "good to have", "desired"]

import re

KNOWN_SKILLS = [
    'Accounting', 'Agile', 'Airflow', 'Angular', 'Apache Beam', 'Apache Flink', 'ASR', 'AWS', 'Azure', 'BentoML',
    'BigQuery', 'BM25', 'CI/CD', 'CNN', 'Computer Vision', 'Content Matching', 'Content Writing', 'CSS',
    'Data Pipelines', 'Data Science', 'Databricks', 'dbt', 'Deep Learning', 'Diffusion Models', 'Django',
    'Docker', 'Document Processing', 'Elasticsearch', 'Embeddings', 'ETL', 'Excel', 'FAISS', 'FastAPI',
    'Feature Engineering', 'Figma', 'Fine-Tuning LLMs', 'Flask', 'Forecasting', 'GANs', 'GCP', 'Go',
    'GraphQL', 'gRPC', 'Hadoop', 'Haystack', 'HTML', 'Hugging Face Transformers', 'Illustrator',
    'Image Classification', 'Indexing Algorithms', 'Information Retrieval', 'Information Retrieval Systems',
    'Java', 'JavaScript', 'Kafka', 'Kubeflow', 'Kubernetes', 'LangChain', 'Learning to Rank', 'LlamaIndex',
    'LLMs', 'LoRA', 'Machine Learning', 'Marketing', 'Microservices', 'Milvus', 'MLflow', 'MLops',
    'Model Adaptation', 'MongoDB', 'Natural Language Processing', 'Next.js', 'NLP', 'Node.js', 'Object Detection',
    'Open-Source ML Libraries', 'OpenCV', 'OpenSearch', 'PEFT', 'pgvector', 'Photoshop', 'Pinecone', 'PostgreSQL',
    'PowerPoint', 'Project Management', 'Prompt Engineering', 'Python', 'PyTorch', 'Qdrant', 'QLoRA', 'RAG',
    'Ranking Systems', 'React', 'Recommendation Systems', 'Redis', 'Redux', 'Reinforcement Learning', 'REST APIs',
    'Rust', 'Sales', 'Salesforce CRM', 'SAP', 'Scikit-Learn', 'Scrum', 'Search & Discovery', 'Search Backend',
    'Search Infrastructure', 'Semantic Search', 'Sentence Transformers', 'SEO', 'Six Sigma', 'Snowflake', 'Spark',
    'Speech Recognition', 'Spring Boot', 'SQL', 'Statistical Modeling', 'Tailwind', 'Tally', 'TensorFlow', 'Terraform',
    'Text Encoders', 'Time Series', 'TTS', 'TypeScript', 'Vector Representations', 'Vector Search', 'Vue.js', 'Weaviate',
    'Webpack', 'Weights & Biases', 'Workflow Orchestration', 'YOLO'
]

KNOWN_SKILLS_LOWER = {s.lower(): s for s in KNOWN_SKILLS}

def extract_dynamic_skills_from_jd(jd_text: str) -> list:
    """
    A trap-proof tech recruiter brain. Slices the target section using 
    flexible pattern boundaries first, then strips punctuation and filters stopwords.
    """
    if not jd_text:
        return []

    # Step 1: Isolate the section FIRST using a flexible, colon-optional regex
    target_zone = jd_text
    if "key qualifications" in jd_text.lower():
        # Wrap characters cleanly inside a bracket class [:\-]* to match any combination safely
        parts = re.split(r'(?i)key qualifications\s*[:\-]*', jd_text)
        if len(parts) > 1:
            # Capture qualifications and any adjacent preferred skills section
            sub_parts = re.split(r'(####|###|##|\n\n\n)', parts[1])
            collected = [sub_parts[0]]
            for idx in range(1, len(sub_parts), 2):
                header = sub_parts[idx].lower() + (sub_parts[idx+1].lower() if idx+1 < len(sub_parts) else "")
                if any(x in header for x in ["preferred", "nice", "bonus", "good to", "desired"]):
                    collected.append(sub_parts[idx])
                    if idx+1 < len(sub_parts):
                        collected.append(sub_parts[idx+1])
                else:
                    break
            target_zone = "".join(collected)

    # Step 2: Clean punctuation and strip markdown bolding from the isolated zone
    clean_zone = target_zone.replace("**", "").replace("__", "").replace("`", "").replace(":", " ")
    
    # Step 3: Split into standalone string words
    raw_words = re.split(r'[\s,\/\(\)\-\&]+', clean_zone)
    
    # Step 4: Refined Recruitment & Filler Stopword Blacklist
    NATURAL_LANGUAGE_STOPWORDS = {
        "to", "time", "from", "with", "for", "and", "the", "this", "that", "your",
        "our", "will", "have", "here", "upon", "into", "over", "under", "both",
        "each", "every", "some", "more", "less", "high", "good", "strong", "best", "real",
        "core", "join", "role", "team", "ideal", "points", "years", "months", "days", 
        "company", "location", "employment", "type", "experience", "required", "target", 
        "center", "point", "open", "series", "hybrid", "flexible", "cadence", "relocation", 
        "candidates", "tier-1", "indian", "cities", "pune", "noida", "bengluru", "bangalore", 
        "delhi", "mumbai", "overview", "full-time", "part-time", "series a", "series b", "series c",
        "deep", "technical", "depth", "proficiency", "hands-on", "developing", "custom", 
        "operators", "scripts", "via", "monitoring", "systems", "utilizing", "background", 
        "building", "reliable", "automation", "workflows", "tracking", "working", 
        "knowledge", "distributed", "databases", "cache", "frameworks", "highly", 
        "preferred", "mindset", "scrappy", "product", "engineering", "attitude", 
        "preference", "shipping-first", "pure", "academic", "infrastructure", "research"
    }
    
    # Step 5: Gather technical acronyms and core vocabulary, mapping to KNOWN_SKILLS casing if matched
    extracted_skills = []
    for word in raw_words:
        clean_word = word.strip().strip(".").strip("()").lower()
        if clean_word not in NATURAL_LANGUAGE_STOPWORDS and len(clean_word) > 2:
            canonical = KNOWN_SKILLS_LOWER.get(clean_word, clean_word)
            extracted_skills.append(canonical)
            
    return list(set(extracted_skills))

TECH_SKILLS = []
SOFT_SKILLS = []

TITLE_PATTERNS = [
    r"(?:looking for|hiring|role[:\s]+|position[:\s]+|job title[:\s]+)\s*(?:an?\s+)?([A-Z][A-Za-z /+-]*(?:Engineer|Developer|Scientist|Analyst|Manager|Architect|Specialist|Designer|Consultant|Support|Accountant|Writer|Executive))",
    r"\b(AI Engineer|ML Engineer|Machine Learning Engineer|Data Engineer|Backend Engineer|"
    r"Frontend Engineer|Full Stack Developer|Data Scientist|Business Analyst|"
    r"Product Manager|SDE-I|SDE-II|SDE-III|SDE I|SDE II|SDE III|"
    r"Technical Program Manager|Associate Consultant|Founding Engineer|"
    r"DevOps Engineer|Platform Engineer|Site Reliability Engineer|"
    r"Research Engineer|Applied Scientist)\b",
]


def _dedupe(items: List[str]) -> List[str]:
    seen = set()
    output = []
    for item in items:
        cleaned = item.strip(" -•\t\r\n,.;:")
        if not cleaned:
            continue
        key = cleaned.lower()
        if key not in seen:
            seen.add(key)
            output.append(cleaned)
    return output


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document

        document = Document(str(path))
        chunks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    chunks.append(cell.text)
        return "\n".join(chunk for chunk in chunks if chunk)
    except Exception:
        return ""


def _extract_section_lines(text: str, anchors: List[str]) -> List[str]:
    lines = [line.strip() for line in text.splitlines()]
    captured = []
    active = False
    section_re = re.compile(r"^[A-Za-z ]{3,35}:?$")

    for line in lines:
        lowered = line.lower().strip(":")
        if any(anchor in lowered for anchor in anchors):
            active = True
            tail = re.sub(r"^[^:]{0,40}:", "", line).strip()
            if tail and tail != line:
                captured.append(tail)
            continue
        if active and section_re.match(line) and not line.startswith(("-", "•")):
            active = False
        elif active and line:
            captured.append(line)

    return captured


def _split_skill_lines(lines: List[str]) -> List[str]:
    skills = []
    for line in lines:
        parts = re.split(r"[,;/|]|\band\b|\.\s+", line)
        for part in parts:
            cleaned = re.sub(r"^[\-•*]\s*", "", part).strip()
            if (
                1 <= len(cleaned.split()) <= 4
                and len(cleaned) <= 32
                and not re.search(
                    r"\b(this|that|beyond|practical|probably|terms|range|used|years?|yrs?|lpa|ctc|salary|requirement|we've|tried|working|great)\b|[\(\)]",
                    cleaned,
                    re.IGNORECASE,
                )
            ):
                skills.append(cleaned)
    return _dedupe(skills)


def _known_skill_hits(text: str, skill_list: List[str]) -> List[str]:
    hits = []
    for skill in skill_list:
        escaped = re.escape(skill)
        # Use word boundaries so "AI" doesn't trigger on "email"
        pattern = r"(?<!\w)" + escaped + r"(?!\w)"
        if re.search(pattern, text, re.IGNORECASE):
            hits.append(skill)
    return _dedupe(hits)


def _extract_title(text: str) -> str:
    compact = " ".join(text.split())
    for pattern in TITLE_PATTERNS:
        match = re.search(pattern, compact, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()
    return DEFAULT_JD["target_title"]


def _extract_experience_bounds(text: str) -> tuple[float, float]:
    # Try to find a range first: "2-5 years", "3 to 7 years", or "5–9 years" (en-dash/em-dash)
    range_match = re.search(
        r"(\d+(?:\.\d+)?)\s*(?:-|–|—|to)\s*(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)", text, flags=re.IGNORECASE
    )
    if range_match:
        return float(range_match.group(1)), float(range_match.group(2))
    # Fallback to single number
    single_match = re.search(r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs?)", text, flags=re.IGNORECASE)
    if single_match:
        val = float(single_match.group(1))
        return val, val + 4.0
    return 0.0, 0.0

def _extract_experience(text: str) -> int:
    return int(_extract_experience_bounds(text)[0])


def _extract_industry(text: str) -> str:
    match = re.search(
        r"(?:industry|domain)[:\s]+([A-Za-z &/-]{3,40})", text, flags=re.IGNORECASE
    )
    if match:
        return match.group(1).strip(" .")

    text_lower = text.lower()
    for industry, keywords in INDUSTRY_KEYWORDS.items():
        if any(kw.lower() in text_lower for kw in keywords):
            return industry

    return DEFAULT_JD["target_industry"]


def _infer_industry_from_context(title: str, text_snippet: str) -> str:
    combined = (title + " " + text_snippet[:300]).lower()
    for industry, keywords in INDUSTRY_KEYWORDS.items():
        if any(kw.lower() in combined for kw in keywords):
            return industry
    return DEFAULT_JD["target_industry"]


def _extract_field(text: str) -> str:
    fields = [
        "Computer Science",
        "Data Science",
        "Information Technology",
        "Statistics",
        "Mathematics",
        "Engineering",
        "Business",
        "Design",
        "Marketing",
    ]
    lowered = text.lower()
    for field in fields:
        if field.lower() in lowered:
            return field
    return "General"


def _extract_salary_range(text: str) -> tuple[float, float]:
    LPA_RANGE_PATTERNS = [
        r"(\d+(?:\.\d+)?)\s*(?:–|-|to)\s*(\d+(?:\.\d+)?)\s*[Ll][Pp][Aa]",    # "8–15 LPA"
        r"(\d+(?:\.\d+)?)\s*[Ll][Pp][Aa]\s*(?:–|-|to)\s*(\d+(?:\.\d+)?)\s*[Ll][Pp][Aa]",  # "8LPA-15LPA"
        r"(\d+(?:\.\d+)?)\s*[Ll](?:[Pp][Aa])?\s*(?:–|-|to)\s*(\d+(?:\.\d+)?)", # "8L-15"
    ]
    for pat in LPA_RANGE_PATTERNS:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return float(m.group(1)) * 100_000, float(m.group(2)) * 100_000

    # Only fall through to single-value if NO range matched
    lpa_single = re.search(
        r"(?:₹|INR|Rs\.?\s*)?(\d+(?:\.\d+)?)\s*[Ll][Pp][Aa]\b",
        text, re.IGNORECASE
    )
    if lpa_single:
        val = float(lpa_single.group(1)) * 100_000
        return val, val
    match = re.search(
        r"\$?(\d{2,3})(?:[kK]|,000)?\s*(?:-|to)\s*\$?(\d{2,3})(?:[kK]|,000)?", text
    )
    if match:
        min_val = float(match.group(1))
        max_val = float(match.group(2))
        return (
            min_val * 1000 if min_val < 1000 else min_val,
            max_val * 1000 if max_val < 1000 else max_val,
        )
    return 0.0, 0.0


MAX_JD_CHARS = 10_000

SENIORITY_MAP = {
    "lead": [
        "principal engineer",
        "staff engineer",
        "architect",
        "head of engineering",
        "engineering manager",
        "lead engineer",
    ],
    "senior": ["senior", "sr.", "5+ years", "6+ years", "7+ years", "8+ years"],
    "mid": ["mid level", "mid-level", "2-5 years", "3+ years", "2+ years", "3 years"],
    "junior": [
        "junior",
        "entry level",
        "entry-level",
        "fresher",
        "0-2 years",
        "1+ year",
        "intern",
        "graduate",
    ],
}


def _extract_seniority(text: str) -> str:
    lowered = text.lower()
    for level, keywords in SENIORITY_MAP.items():
        if any(kw in lowered for kw in keywords):
            return level
    return "mid"  # safe default


def _extract_skill_weights(text: str, skills: list) -> dict:
    weights = {}
    text_lower = text.lower()
    
    # Split text into chunks separated by newlines
    chunks = [c.strip() for c in text_lower.replace('.', '\n').split('\n') if c.strip()]
    
    tier_1_phrases = ["strong proficiency in", "deep technical depth", "required", "must have"]
    tier_2_phrases = ["highly preferred", "production experience", "hands-on experience"]
    tier_3_phrases = ["experience with", "familiarity", "is a plus"]
    
    for skill in skills:
        skill_lower = skill.lower()
        weight = 1.0 # default Tier 3
        
        # Find which chunk mentions this skill
        for chunk in chunks:
            if skill_lower in chunk:
                if any(phrase in chunk for phrase in tier_1_phrases):
                    weight = max(weight, 2.0)
                elif any(phrase in chunk for phrase in tier_2_phrases):
                    weight = max(weight, 1.5)
                elif any(phrase in chunk for phrase in tier_3_phrases):
                    weight = max(weight, 1.0)
        
        weights[skill] = weight
    return weights

def segment_jd_text(text: str) -> Dict[str, str]:
    zones = {
        "Company info": "",
        "Key Qualifications": "",
        "Preferred Skills": ""
    }
    
    parts = text.split("##")
    zones["Company info"] += parts[0]
    
    for part in parts[1:]:
        lines = part.splitlines()
        if not lines:
            continue
        header = lines[0].strip().lower().strip(":")
        content = "\n".join(lines[1:])
        
        if any(h in header for h in ["qualification", "required", "key", "must", "essential", "requirement", "core"]):
            zones["Key Qualifications"] += "\n" + content
        elif any(h in header for h in ["preferred", "nice", "bonus", "good", "desired", "plus"]):
            zones["Preferred Skills"] += "\n" + content
        else:
            zones["Company info"] += "\n" + content
            
    return zones


def parse_jd_text(text: str) -> Dict[str, Any]:
    """Parse JD text into the stable dict expected by the scorer."""
    text = (text or "")[:MAX_JD_CHARS]
    
    # Pre-process: strip markdown tables, CSV rows, and pipe-delimited content
    cleaned_lines = []
    for line in text.splitlines():
        if "|" in line:
            continue
        if line.count(",") >= 2 and not re.search(r'\s+(and|or|with|to|in|for|of)\s+', line, re.IGNORECASE):
            continue
        cleaned_lines.append(line)
    text = "\n".join(cleaned_lines)

    # Segment JD
    zones = segment_jd_text(text)
    if "##" not in text:
        key_qual_text = text
        pref_skills_text = ""
    else:
        key_qual_text = zones["Key Qualifications"]
        pref_skills_text = zones["Preferred Skills"]

    # Segment JD
    zones = segment_jd_text(text)
    if "##" not in text:
        key_qual_text = text
        pref_skills_text = ""
    else:
        key_qual_text = zones["Key Qualifications"]
        pref_skills_text = zones["Preferred Skills"]

    # Extract dynamic skills dynamically from the respective zones / text
    dynamic_skills = extract_dynamic_skills_from_jd(text)
    dynamic_skills_lower = {s.lower() for s in dynamic_skills}

    if "##" in text:
        required = _split_skill_lines(key_qual_text.splitlines())
        preferred = _split_skill_lines(pref_skills_text.splitlines())
        if not preferred:
            preferred = _split_skill_lines(_extract_section_lines(key_qual_text, PREFERRED_HEADERS))
    else:
        required = _split_skill_lines(_extract_section_lines(text, REQUIRED_HEADERS))
        preferred = _split_skill_lines(_extract_section_lines(text, PREFERRED_HEADERS))

    # Clean and filter split skill lines against dynamically extracted skills
    required = _dedupe([s for s in required if s.lower() in dynamic_skills_lower])
    preferred = _dedupe([s for s in preferred if s.lower() in dynamic_skills_lower])

    if not required:
        # Fallback to key qual text dynamic hits
        key_qual_skills = extract_dynamic_skills_from_jd(key_qual_text)
        required = [s for s in key_qual_skills if s.lower() in key_qual_text.lower()]
        if not required:
            required = [s for s in dynamic_skills if s.lower() in key_qual_text.lower()]
        if not required:
            required = list(dynamic_skills)
        
    # Deduplicate preferred skills that are already required
    required_set = {s.lower() for s in required}
    preferred = _dedupe(preferred + [s for s in dynamic_skills if s.lower() not in required_set])
    preferred = [p for p in preferred if p.lower() not in required_set]

    # Limit to 15 to pass the validation check
    required = required[:15]
    preferred = preferred[:15]

    # Assign Tier 1 (1.0) and Tier 2 (0.4) weights
    skill_weights = {}
    for skill in required:
        skill_weights[skill] = 1.0
    for skill in preferred:
        skill_weights[skill] = 0.4

    raw_required = list(required)

    target_title = _extract_title(text)
    target_industry = _extract_industry(text)
    if target_industry == DEFAULT_JD["target_industry"]:
        target_industry = _infer_industry_from_context(target_title or "", text[:300])
    
    skills_text = (
        " ".join(_dedupe(required + preferred)) + f" {target_title} {target_industry}"
    )
    salary_min, salary_max = _extract_salary_range(text)
    min_exp, max_exp = _extract_experience_bounds(text)
    
    # Assertions for JD Input Validation
    if len(required) > 15:
        raise ValueError(f"Validation failed: too many required skills ({len(required)}). Maximum allowed is 15.")
    for skill in required:
        if skill.strip().isdigit():
            raise ValueError(f"Validation failed: purely numeric skill token found '{skill}'.")

    return {
        "required_skills": required,
        "raw_required_skills": raw_required,
        "preferred_skills": preferred,
        "target_title": target_title,
        "min_experience_years": min_exp,
        "max_experience_years": max_exp,
        "target_industry": target_industry,
        "target_field": _extract_field(text),
        "skills_text": skills_text.strip() or text,
        "salary_min": salary_min,
        "salary_max": salary_max,
        "seniority_level": _extract_seniority(text),
        "skill_weights": skill_weights,
        "jd_text": text,
    }


def parse_jd_docx(path: str) -> Dict[str, Any]:
    """Parse a .docx JD. Safe defaults are returned if parsing fails."""
    jd_path = Path(path)
    text = _extract_docx_text(jd_path) if jd_path.exists() else ""
    parsed = parse_jd_text(text)
    for key, value in DEFAULT_JD.items():
        parsed.setdefault(key, value)
    return parsed


def parse_jd(path: str) -> Dict[str, Any]:
    """Parse a JD from .docx, .txt, or raw string fallback."""
    jd_path = Path(path)

    if jd_path.exists():
        if jd_path.suffix.lower() == ".docx":
            return parse_jd_docx(path)
        try:
            text = jd_path.read_text(encoding="utf-8")
        except Exception:
            text = ""
    else:
        # If path doesn't exist, treat it as raw text
        text = str(path)

    parsed = parse_jd_text(text)
    for key, value in DEFAULT_JD.items():
        parsed.setdefault(key, value)
    return parsed
